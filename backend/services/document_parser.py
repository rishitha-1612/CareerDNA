import io
import re
from pathlib import Path

def extract_text_from_pdf(file_bytes: bytes) -> str:
    from pdfminer.high_level import extract_text_to_fp
    from pdfminer.layout import LAParams
    output = io.StringIO()
    extract_text_to_fp(io.BytesIO(file_bytes), output, laparams=LAParams(), output_type="text", codec="utf-8")
    return _clean_text(output.getvalue())

def extract_text_from_docx(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return _clean_text("\n".join(paragraphs))

def _clean_text(raw: str) -> str:
    text = re.sub(r"\s+", " ", raw)
    text = re.sub(r"[^\x20-\x7E\n]", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

def detect_and_extract(file_bytes: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_bytes)
    elif ext in (".txt", ".md"):
        return _clean_text(file_bytes.decode("utf-8", errors="ignore"))
    else:
        raise ValueError(f"Unsupported file type: {ext}. Use PDF, DOCX, or TXT.")
